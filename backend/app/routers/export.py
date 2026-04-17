import io
import json
from datetime import datetime

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from ..database import get_db, Script, Segment, SegmentAnalysis, ScriptReport
from ..schemas import ExportRequest

router = APIRouter(prefix="/api/scripts", tags=["export"])


def _build_json_export(script: Script, segments: list, analyses: dict, report) -> dict:
    segs = []
    for seg in segments:
        a = analyses.get(seg.id)
        segs.append({
            "index": seg.index,
            "raw_text": seg.raw_text,
            "start_char": seg.start_char,
            "end_char": seg.end_char,
            "structure": a.structure if a else "",
            "attitude_type": a.attitude_type if a else "",
            "attitude_object": a.attitude_object if a else "",
            "attitude_insight": a.attitude_insight if a else "",
            "techniques": a.techniques if a else "",
            "problems": a.problems if a else "",
            "notes": a.notes if a else "",
            "inspiration": a.inspiration if a else "",
            "analysis_text": a.analysis_text if a else "",
            "starred": bool(a.starred) if a else False,
        })

    return {
        "title": script.title or script.filename,
        "actor": script.actor_name or "",
        "show": script.show_name or "",
        "exported_at": datetime.utcnow().isoformat(),
        "report": {
            "summary": report.summary if report else "",
            "strengths": report.strengths if report else "",
            "weaknesses": report.weaknesses if report else "",
            "methodology": report.methodology if report else "",
            "key_insights": report.key_insights if report else "",
            "overall_score": report.overall_score if report else 0.0,
        } if report else None,
        "segments": segs,
    }


def _build_md_export(script: Script, segments: list, analyses: dict, report) -> str:
    lines = []
    lines.append(f"# {script.title or script.filename}")
    if script.actor_name:
        lines.append(f"**演员**: {script.actor_name}")
    if script.show_name:
        lines.append(f"**节目**: {script.show_name}")
    lines.append("")

    if report:
        lines.append("## 整篇总结")
        if report.summary:
            lines.append(report.summary)
        if report.strengths:
            lines.append("\n### 强项\n" + report.strengths)
        if report.weaknesses:
            lines.append("\n### 弱项\n" + report.weaknesses)
        if report.methodology:
            lines.append("\n### 方法论\n" + report.methodology)
        if report.key_insights:
            lines.append("\n### 关键洞察\n" + report.key_insights)
        if report.overall_score:
            lines.append(f"\n**评分**: {report.overall_score:.2f}")
        lines.append("")

    lines.append(f"## 段落分析 ({len(segments)}段)")
    lines.append("")

    for seg in segments:
        a = analyses.get(seg.id)
        lines.append(f"### 段落 {seg.index + 1}")
        lines.append(f"> {seg.raw_text}")
        lines.append("")
        if a:
            if a.structure:
                lines.append(f"- **结构**: {a.structure} — {a.structure_note}")
            if a.attitude_type:
                lines.append(f"- **态度**: {a.attitude_type} ({a.attitude_object}) — {a.attitude_insight}")
            if a.techniques:
                lines.append(f"- **技巧**: {a.techniques} — {a.technique_notes}")
            if a.problems:
                lines.append(f"- **问题**: {a.problems} — {a.problem_notes}")
            if a.notes:
                lines.append(f"- **备注**: {a.notes}")
            if a.inspiration:
                lines.append(f"- **启发**: {a.inspiration}")
            if a.analysis_text:
                lines.append(f"\n{a.analysis_text}")
            if a.starred:
                lines.append("\n⭐ 已收藏")
        lines.append("")
        lines.append("---")
        lines.append("")

    return "\n".join(lines)


def _build_docx_export(script: Script, segments: list, analyses: dict, report) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_heading(script.title or script.filename, 0)
    if script.actor_name:
        doc.add_paragraph(f"演员: {script.actor_name}")
    if script.show_name:
        doc.add_paragraph(f"节目: {script.show_name}")

    if report:
        doc.add_heading("整篇总结", 1)
        if report.summary:
            doc.add_paragraph(report.summary)
        if report.strengths:
            doc.add_heading("强项", 2)
            doc.add_paragraph(report.strengths)
        if report.weaknesses:
            doc.add_heading("弱项", 2)
            doc.add_paragraph(report.weaknesses)
        if report.methodology:
            doc.add_heading("方法论", 2)
            doc.add_paragraph(report.methodology)
        if report.key_insights:
            doc.add_heading("关键洞察", 2)
            doc.add_paragraph(report.key_insights)
        if report.overall_score:
            doc.add_paragraph(f"评分: {report.overall_score:.2f}")

    doc.add_heading("段落分析", 1)
    for seg in segments:
        a = analyses.get(seg.id)
        doc.add_heading(f"段落 {seg.index + 1}", 2)
        p = doc.add_paragraph()
        run = p.add_run(seg.raw_text)
        run.font.italic = True
        if a:
            if a.structure:
                doc.add_paragraph(f"结构: {a.structure} — {a.structure_note}")
            if a.attitude_type:
                doc.add_paragraph(f"态度: {a.attitude_type} ({a.attitude_object}) — {a.attitude_insight}")
            if a.techniques:
                doc.add_paragraph(f"技巧: {a.techniques} — {a.technique_notes}")
            if a.problems:
                doc.add_paragraph(f"问题: {a.problems} — {a.problem_notes}")
            if a.notes:
                doc.add_paragraph(f"备注: {a.notes}")
            if a.inspiration:
                doc.add_paragraph(f"启发: {a.inspiration}")
            if a.analysis_text:
                doc.add_paragraph(a.analysis_text)
            if a.starred:
                doc.add_paragraph("⭐ 已收藏")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@router.post("/{script_id}/export")
def export_script_post(
    script_id: int,
    data: ExportRequest,
    db: Session = Depends(get_db),
):
    return _do_export(script_id, data.format, data.include_raw, data.include_analysis, db)


@router.get("/{script_id}/export")
def export_script_get(
    script_id: int,
    format: str = "json",
    include_raw: bool = True,
    include_analysis: bool = True,
    db: Session = Depends(get_db),
):
    return _do_export(script_id, format, include_raw, include_analysis, db)


def _do_export(script_id: int, fmt: str, include_raw: bool, include_analysis: bool, db: Session):
    script = db.query(Script).filter(Script.id == script_id).first()
    if not script:
        raise HTTPException(404, "Script not found")

    segments = db.query(Segment).filter(Segment.script_id == script_id).order_by(Segment.index).all()
    analyses = {}
    for seg in segments:
        a = db.query(SegmentAnalysis).filter(SegmentAnalysis.segment_id == seg.id).first()
        if a:
            analyses[seg.id] = a

    report = db.query(ScriptReport).filter(ScriptReport.script_id == script_id).first()

    filename_base = (script.title or script.filename or "export").replace(" ", "_")
    import urllib.parse
    safe_name = urllib.parse.quote(filename_base)

    if fmt == "json":
        payload = _build_json_export(script, segments, analyses, report)
        content = json.dumps(payload, ensure_ascii=False, indent=2)
        return StreamingResponse(
            io.StringIO(content),
            media_type="application/json",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}.json"; filename*=UTF-8\'\'{safe_name}.json'},
        )

    elif fmt == "md":
        content = _build_md_export(script, segments, analyses, report)
        return StreamingResponse(
            io.StringIO(content),
            media_type="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}.md"; filename*=UTF-8\'\'{safe_name}.md'},
        )

    elif fmt == "docx":
        content = _build_docx_export(script, segments, analyses, report)
        return StreamingResponse(
            io.BytesIO(content),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}.docx"; filename*=UTF-8\'\'{safe_name}.docx'},
        )

    raise HTTPException(400, "Invalid format")
